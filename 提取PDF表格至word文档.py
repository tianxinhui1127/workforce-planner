import os
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_content_from_pdf(pdf_path):
    """提取PDF中的文本和表格"""
    all_content = []
    
    try:
        # 使用pdfplumber提取内容
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                
                
                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    # 过滤空表格
                    if any(any(cell for cell in row if cell) for row in table):
                        all_content.append({"type": "table", "content": table})
                
                # 添加页码标记
                all_content.append({"type": "page_break"})
        
        return all_content
    except Exception as e:
        print(f"提取内容时发生错误: {str(e)}")
        return []

def save_to_word(content, output_path):
    """将内容保存到Word文档"""
    try:
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
        
        # 设置页面边距
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        for item in content:
            if item["type"] == "text":
                # ... existing text handling code ...
                para = doc.add_paragraph(item["content"])
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
            elif item["type"] == "table":
                table_data = item["content"]
                if table_data:
                    # 获取行数和列数
                    rows = len(table_data)
                    cols = len(table_data[0])
                    
                    # 创建表格
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    table.autofit = False  # 关闭自动适应单元格内容
                    
                    # 填充表格内容
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            cell_text = str(cell) if cell is not None else ""
                            table.cell(i, j).text = cell_text.strip()
                    
                    # 设置表格宽度为页面宽度
                    table.allow_autofit = True
                    for row in table.rows:
                        for cell in row.cells:
                            # 设置单元格自动换行
                            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.paragraph_format.space_before = Pt(0)
                    
                    # 设置每列宽度相等
                    for column in table.columns:
                        column.width = Inches(6.5 / cols)  # 平均分配页面宽度
                
                # 在表格后添加空行
                doc.add_paragraph()
                
            elif item["type"] == "page_break":
                doc.add_page_break()
        
        # 保存文档
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"保存Word文档时发生错误: {str(e)}")
        return False

def main():
    try:
        # 指定PDF文件路径
        pdf_path = "C:/Users/市场开发技术组/Desktop/页面从 1标资审文件.pdf"
        
        # 获取PDF文件所在目录和文件名（不含扩展名）
        pdf_dir = os.path.dirname(pdf_path)
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        # 构建输出文件路径
        word_output_path = os.path.join(pdf_dir, f"{pdf_name}_导出.docx")
        
        print("正在提取PDF内容...")
        content = extract_content_from_pdf(pdf_path)
        
        if content:
            print("正在生成Word文档...")
            if save_to_word(content, word_output_path):
                print(f"转换完成！文件已保存到: {word_output_path}")
            else:
                print("保存Word文档失败")
        else:
            print("未能提取到内容")
            
    except Exception as e:
        print(f"发生错误: {str(e)}")

if __name__ == "__main__":
    main()

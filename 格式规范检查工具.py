import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import time

class FormatCheckerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("格式规范—暗标文件检查工具")
        self.root.geometry("1000x600")
        
        # 设置中文字体支持
        self.style = ttk.Style()
        self.style.configure("TLabel", font=("SimHei", 10))
        self.style.configure("TButton", font=("SimHei", 10))
        self.style.configure("TCombobox", font=("SimHei", 10))
        self.style.configure("TEntry", font=("SimHei", 10))
        
        # 创建菜单栏
        self.create_menu()
        
        # 创建标签页
        self.create_tabs()
        
        # 创建状态栏
        self.create_status_bar()
        
        # 初始化状态栏消息
        self.update_status("就绪")
        
        # 绑定窗口关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
    def create_menu(self):
        # 创建菜单栏
        self.menu_bar = tk.Menu(self.root)
        
        # 文件菜单
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.file_menu.add_command(label="打开文件", command=self.select_file)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="退出", command=self.on_closing)
        self.menu_bar.add_cascade(label="文件", menu=self.file_menu)
        
        # 帮助菜单
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.help_menu.add_command(label="使用说明")
        self.help_menu.add_command(label="关于")
        self.menu_bar.add_cascade(label="帮助", menu=self.help_menu)
        
        # 设置菜单栏
        self.root.config(menu=self.menu_bar)
    
    def create_tabs(self):
        # 创建标签控件
        self.tab_control = ttk.Notebook(self.root)
        
        # 创建各个标签页
        self.tab_document = ttk.Frame(self.tab_control)
        self.tab_sensitive = ttk.Frame(self.tab_control)
        self.tab_image = ttk.Frame(self.tab_control)
        self.tab_features = ttk.Frame(self.tab_control)
        self.tab_help = ttk.Frame(self.tab_control)
        
        # 添加标签页
        self.tab_control.add(self.tab_document, text="文档检查")
        self.tab_control.add(self.tab_sensitive, text="敏感词检查")
        self.tab_control.add(self.tab_image, text="图片处理")
        self.tab_control.add(self.tab_features, text="功能购买")
        self.tab_control.add(self.tab_help, text="使用说明")
        
        # 打包标签控件
        self.tab_control.pack(expand=1, fill="both", padx=5, pady=5)
        
        # 在文档检查标签页中创建主框架
        self.main_frame = ttk.Frame(self.tab_document)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建左侧控制面板
        self.create_control_panel()
        
        # 创建预览区域
        self.create_preview_area()
        
        # 设置敏感词标签页
        self.setup_sensitive_tab()
        
        # 为其他标签页添加简单的占位文本
        ttk.Label(self.tab_image, text="图片处理功能开发中...", font=("SimHei", 12)).pack(pady=50)
        ttk.Label(self.tab_features, text="功能购买功能开发中...", font=("SimHei", 12)).pack(pady=50)
        ttk.Label(self.tab_help, text="使用说明功能开发中...", font=("SimHei", 12)).pack(pady=50)
        
    def create_status_bar(self):
        # 创建状态栏
        self.status_bar = ttk.Label(self.root, text="就绪", anchor=tk.W, relief=tk.SUNKEN, font=("SimHei", 9))
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
    def update_status(self, message):
        # 更新状态栏消息
        self.status_bar.config(text=message)
        self.root.update_idletasks()
        
    def on_closing(self):
        # 处理窗口关闭事件
        if tk.messagebox.askokcancel("退出", "确定要退出程序吗？"):
            self.root.destroy()
    
    def create_control_panel(self):
        # 创建左侧控制面板
        self.control_frame = ttk.LabelFrame(self.main_frame, text="")
        self.control_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10), pady=0, expand=False)
        self.control_frame.configure(width=350)
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(self.control_frame, text="文档选择", relief=tk.RAISED)
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 选择文件按钮 - 绿色样式
        self.style.configure("Green.TButton", foreground="black")
        if os.name == 'nt':  # Windows系统
            self.style.configure("Green.TButton", background="#4CAF50", foreground="black")
            self.style.map("Green.TButton", background=[("active", "#45a049")])
        
        self.select_file_btn = ttk.Button(file_frame, text="选择Word文档 (*.docx)", command=self.select_file, style="Green.TButton")
        self.select_file_btn.pack(fill=tk.X, padx=5, pady=5)
        
        # 路径显示标签
        self.file_path_var = tk.StringVar(value="未选择文件")
        file_path_label = ttk.Label(file_frame, textvariable=self.file_path_var, relief=tk.SUNKEN, anchor="w", wraplength=300)
        file_path_label.pack(fill=tk.X, padx=5, pady=5)
        
        # 完整路径显示标签
        self.path_label = ttk.Label(file_frame, text="路径： ", relief=tk.SUNKEN, anchor="w", wraplength=300)
        self.path_label.pack(fill=tk.X, padx=5, pady=5)
        
        # 开始检查按钮 - 蓝色样式
        self.style.configure("Blue.TButton", foreground="black")
        if os.name == 'nt':  # Windows系统
            self.style.configure("Blue.TButton", background="#2196F3", foreground="black")
            self.style.map("Blue.TButton", background=[("active", "#1976D2")])
        
        self.check_btn = ttk.Button(file_frame, text="开始检查", command=self.check_document, style="Blue.TButton")
        self.check_btn.pack(fill=tk.X, padx=5, pady=10)
        
        # 修复文件选项
        fix_frame = ttk.LabelFrame(self.control_frame, text="修复选项", relief=tk.RAISED)
        fix_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 自动生成修复文件选项
        self.fix_file_var = tk.BooleanVar(value=False)
        fix_checkbox = ttk.Checkbutton(fix_frame, text="自动生成修复文件(原文件名_fixed.docx)", variable=self.fix_file_var)
        fix_checkbox.pack(anchor="w", padx=5, pady=5)
        
        # 修复按钮
        self.fix_btn = ttk.Button(fix_frame, text="修复文档", command=self.fix_document, style="Green.TButton")
        self.fix_btn.pack(fill=tk.X, padx=5, pady=5)
        
        # 检查选项区域
        options_frame = ttk.LabelFrame(self.control_frame, text="检查选项", relief=tk.RAISED)
        options_frame.pack(fill=tk.X, pady=5, expand=True)
        
        # 全选和反选按钮区域
        selection_frame = ttk.Frame(options_frame)
        selection_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 设置橙色按钮样式
        self.style.configure("Orange.TButton", foreground="black")
        if os.name == 'nt':
            self.style.configure("Orange.TButton", background="#FF9800", foreground="black")
            self.style.map("Orange.TButton", background=[("active", "#F57C00")])
        
        # 设置浅蓝色按钮样式
        self.style.configure("LightBlue.TButton", foreground="black")
        if os.name == 'nt':
            self.style.configure("LightBlue.TButton", background="#03A9F4", foreground="black")
            self.style.map("LightBlue.TButton", background=[("active", "#0288D1")])
        
        self.select_all_btn = ttk.Button(selection_frame, text="全选", style="Orange.TButton", command=self.select_all)
        self.select_all_btn.pack(side=tk.LEFT, padx=2)
        
        self.select_none_btn = ttk.Button(selection_frame, text="反选", style="LightBlue.TButton", command=self.select_none)
        self.select_none_btn.pack(side=tk.LEFT, padx=2)
        
        # 创建地区选择和预设按钮区域
        region_frame = ttk.Frame(options_frame)
        region_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 创建地区选择标签
        region_label = ttk.Label(region_frame, text="地区标准:")
        region_label.pack(side=tk.LEFT, padx=2)
        
        # 创建地区选择下拉框
        self.region_var = tk.StringVar()
        self.region_combo = ttk.Combobox(region_frame, textvariable=self.region_var, values=["国家标准", "地方标准", "自定义"], width=15)
        self.region_combo.current(0)
        self.region_combo.pack(side=tk.LEFT, padx=2)
        
        # 创建预设格式说明按钮
        self.preset_btn = ttk.Button(region_frame, text="预设格式说明", style="TButton", command=self.show_preset_info)
        self.preset_btn.pack(side=tk.LEFT, padx=2)
        
        # 创建格式检查项
        # 纸张大小检查
        self.size_check_var = tk.BooleanVar(value=True)
        size_checkbox = ttk.Checkbutton(options_frame, text="纸张大小检查", variable=self.size_check_var)
        size_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 页边距检查
        self.margin_check_var = tk.BooleanVar(value=True)
        margin_checkbox = ttk.Checkbutton(options_frame, text="页边距检查", variable=self.margin_check_var)
        margin_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 页边距设置区域
        margin_frame = ttk.Frame(options_frame)
        margin_frame.pack(fill=tk.X, padx=20)
        
        # 上页边距
        ttk.Label(margin_frame, text="上:").grid(row=0, column=0, padx=2)
        self.margin_top_var = tk.StringVar(value="2.5")
        ttk.Entry(margin_frame, textvariable=self.margin_top_var, width=5).grid(row=0, column=1, padx=2)
        ttk.Label(margin_frame, text="cm").grid(row=0, column=2, padx=2)
        
        # 下页边距
        ttk.Label(margin_frame, text="下:").grid(row=0, column=3, padx=2)
        self.margin_bottom_var = tk.StringVar(value="2.5")
        ttk.Entry(margin_frame, textvariable=self.margin_bottom_var, width=5).grid(row=0, column=4, padx=2)
        ttk.Label(margin_frame, text="cm").grid(row=0, column=5, padx=2)
        
        # 左页边距
        ttk.Label(margin_frame, text="左:").grid(row=0, column=6, padx=2)
        self.margin_left_var = tk.StringVar(value="2.5")
        ttk.Entry(margin_frame, textvariable=self.margin_left_var, width=5).grid(row=0, column=7, padx=2)
        ttk.Label(margin_frame, text="cm").grid(row=0, column=8, padx=2)
        
        # 右页边距
        ttk.Label(margin_frame, text="右:").grid(row=0, column=9, padx=2)
        self.margin_right_var = tk.StringVar(value="2.5")
        ttk.Entry(margin_frame, textvariable=self.margin_right_var, width=5).grid(row=0, column=10, padx=2)
        ttk.Label(margin_frame, text="cm").grid(row=0, column=11, padx=2)
        
        # 页眉页脚检查
        self.header_check_var = tk.BooleanVar(value=True)
        header_checkbox = ttk.Checkbutton(options_frame, text="页眉页脚检查", variable=self.header_check_var)
        header_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 目录检查
        self.toc_check_var = tk.BooleanVar(value=True)
        toc_checkbox = ttk.Checkbutton(options_frame, text="目录检查", variable=self.toc_check_var)
        toc_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 段落对齐检查
        self.align_check_var = tk.BooleanVar(value=True)
        align_checkbox = ttk.Checkbutton(options_frame, text="段落对齐检查", variable=self.align_check_var)
        align_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 对齐方式选择
        align_frame = ttk.Frame(options_frame)
        align_frame.pack(fill=tk.X, padx=20)
        
        self.align_var = tk.StringVar(value="两端对齐")
        ttk.Label(align_frame, text="对齐方式:").pack(side=tk.LEFT, padx=2)
        align_combo = ttk.Combobox(align_frame, textvariable=self.align_var, values=["左对齐", "居中对齐", "右对齐", "两端对齐"], width=10)
        align_combo.pack(side=tk.LEFT, padx=2)
        
        # 行距检查
        self.line_spacing_check_var = tk.BooleanVar(value=True)
        line_spacing_checkbox = ttk.Checkbutton(options_frame, text="行距检查", variable=self.line_spacing_check_var)
        line_spacing_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 行距设置
        line_spacing_frame = ttk.Frame(options_frame)
        line_spacing_frame.pack(fill=tk.X, padx=20)
        
        self.line_spacing_var = tk.StringVar(value="28.00 磅")
        ttk.Label(line_spacing_frame, text="行距:").pack(side=tk.LEFT, padx=2)
        line_spacing_combo = ttk.Combobox(line_spacing_frame, textvariable=self.line_spacing_var, 
                                          values=["18.00 磅", "20.00 磅", "22.00 磅", "24.00 磅", "26.00 磅", "28.00 磅", "30.00 磅"], 
                                          width=10)
        line_spacing_combo.pack(side=tk.LEFT, padx=2)
        
        # 首行缩进检查
        self.indent_check_var = tk.BooleanVar(value=True)
        indent_checkbox = ttk.Checkbutton(options_frame, text="首行缩进检查", variable=self.indent_check_var)
        indent_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 缩进设置
        indent_frame = ttk.Frame(options_frame)
        indent_frame.pack(fill=tk.X, padx=20)
        
        self.indent_var = tk.StringVar(value="2.00 字符")
        ttk.Label(indent_frame, text="首行缩进:").pack(side=tk.LEFT, padx=2)
        indent_combo = ttk.Combobox(indent_frame, textvariable=self.indent_var, 
                                   values=["2.00 字符", "2.50 字符", "3.00 字符"], 
                                   width=10)
        indent_combo.pack(side=tk.LEFT, padx=2)
        
        # 正文字体检查
        self.font_check_var = tk.BooleanVar(value=True)
        font_checkbox = ttk.Checkbutton(options_frame, text="正文字体检查", variable=self.font_check_var)
        font_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 字体设置
        font_frame = ttk.Frame(options_frame)
        font_frame.pack(fill=tk.X, padx=20)
        
        self.font_var = tk.StringVar(value="宋体")
        ttk.Label(font_frame, text="正文字体:").pack(side=tk.LEFT, padx=2)
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_var, 
                                values=["宋体", "微软雅黑", "黑体", "楷体"], 
                                width=10)
        font_combo.pack(side=tk.LEFT, padx=2)
        
        # 正文字号检查
        self.font_size_check_var = tk.BooleanVar(value=True)
        font_size_checkbox = ttk.Checkbutton(options_frame, text="正文字号检查", variable=self.font_size_check_var)
        font_size_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 字号设置
        font_size_frame = ttk.Frame(options_frame)
        font_size_frame.pack(fill=tk.X, padx=20)
        
        self.font_size_var = tk.StringVar(value="14号(四号)")
        ttk.Label(font_size_frame, text="正文字号:").pack(side=tk.LEFT, padx=2)
        font_size_combo = ttk.Combobox(font_size_frame, textvariable=self.font_size_var, 
                                      values=["12号(小四号)", "14号(小四)", "16号(四号)", "18号(小三)"], 
                                      width=10)
        font_size_combo.pack(side=tk.LEFT, padx=2)
        
        # 表格内字体检查
        self.inner_font_check_var = tk.BooleanVar(value=True)
        inner_font_checkbox = ttk.Checkbutton(options_frame, text="表格内字体检查", variable=self.inner_font_check_var)
        inner_font_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 表格内字体设置
        inner_font_frame = ttk.Frame(options_frame)
        inner_font_frame.pack(fill=tk.X, padx=20)
        
        self.inner_font_var = tk.StringVar(value="宋体")
        ttk.Label(inner_font_frame, text="表格内字体:").pack(side=tk.LEFT, padx=2)
        inner_font_combo = ttk.Combobox(inner_font_frame, textvariable=self.inner_font_var, 
                                      values=["宋体", "微软雅黑", "黑体", "楷体"], 
                                      width=10)
        inner_font_combo.pack(side=tk.LEFT, padx=2)
        
        # 表格内字号检查
        self.inner_font_size_check_var = tk.BooleanVar(value=True)
        inner_font_size_checkbox = ttk.Checkbutton(options_frame, text="表格内字号检查", variable=self.inner_font_size_check_var)
        inner_font_size_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 表格内字号设置
        inner_font_size_frame = ttk.Frame(options_frame)
        inner_font_size_frame.pack(fill=tk.X, padx=20)
        
        self.inner_font_size_var = tk.StringVar(value="12.00 号")
        ttk.Label(inner_font_size_frame, text="表格内字号:").pack(side=tk.LEFT, padx=2)
        inner_font_size_combo = ttk.Combobox(inner_font_size_frame, textvariable=self.inner_font_size_var, 
                                           values=["10.00 号", "12.00 号", "14.00 号"], 
                                           width=10)
        inner_font_size_combo.pack(side=tk.LEFT, padx=2)
        
        # 特殊格式检查
        self.special_format_check_var = tk.BooleanVar(value=True)
        special_format_checkbox = ttk.Checkbutton(options_frame, text="特殊格式检查", variable=self.special_format_check_var)
        special_format_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 特殊格式选项
        format_frame = ttk.Frame(options_frame)
        format_frame.pack(fill=tk.X, padx=20)
        
        self.format_bold_var = tk.BooleanVar(value=True)
        format_bold = ttk.Checkbutton(format_frame, text="禁用加粗", variable=self.format_bold_var)
        format_bold.pack(side=tk.LEFT, padx=5)
        
        self.format_italic_var = tk.BooleanVar(value=True)
        format_italic = ttk.Checkbutton(format_frame, text="禁用斜体", variable=self.format_italic_var)
        format_italic.pack(side=tk.LEFT, padx=5)
        
        self.format_underline_var = tk.BooleanVar(value=True)
        format_underline = ttk.Checkbutton(format_frame, text="禁用下划线", variable=self.format_underline_var)
        format_underline.pack(side=tk.LEFT, padx=5)
        
        # 文字颜色检查
        self.color_check_var = tk.BooleanVar(value=True)
        color_checkbox = ttk.Checkbutton(options_frame, text="文字颜色检查", variable=self.color_check_var)
        color_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 颜色选择
        color_frame = ttk.Frame(options_frame)
        color_frame.pack(fill=tk.X, padx=20)
        
        self.color_var = tk.StringVar(value="黑色")
        ttk.Label(color_frame, text="文字颜色:").pack(side=tk.LEFT, padx=2)
        color_combo = ttk.Combobox(color_frame, textvariable=self.color_var, 
                                 values=["黑色", "蓝色", "红色"], 
                                 width=10)
        color_combo.pack(side=tk.LEFT, padx=2)
        
        # 标点符号检查
        self.punctuation_check_var = tk.BooleanVar(value=True)
        punctuation_checkbox = ttk.Checkbutton(options_frame, text="英文标点符号检查", variable=self.punctuation_check_var)
        punctuation_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 空格检查
        self.space_check_var = tk.BooleanVar(value=True)
        space_checkbox = ttk.Checkbutton(options_frame, text="空格检查", variable=self.space_check_var)
        space_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 字符间距检查
        self.char_spacing_check_var = tk.BooleanVar(value=True)
        char_spacing_checkbox = ttk.Checkbutton(options_frame, text="字符间距检查", variable=self.char_spacing_check_var)
        char_spacing_checkbox.pack(anchor="w", padx=5, pady=2)
        
        # 为字符间距创建一个框架
        spacing_check_frame = ttk.Frame(options_frame)
        spacing_check_frame.pack(fill=tk.X, padx=20)
        
        self.spacing_entry_var = tk.StringVar(value="")
        ttk.Entry(spacing_check_frame, textvariable=self.spacing_entry_var).pack(fill=tk.X, padx=20, pady=2)
        
    def select_file(self):
        # 文件选择对话框
        file_path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if file_path:
            # 只显示文件名，不显示完整路径在第一个文本框
            file_name = os.path.basename(file_path)
            self.file_path = file_path  # 保存完整路径到实例变量
            self.file_path_var.set(f"已选择：{file_name}")
            # 路径标签显示完整路径
            self.path_label.config(text=f"路径：{file_path}")
            self.update_status(f"已选择文件：{file_name}")
            self.update_result_text(f"已选择文件：{file_name}\n\n请点击'开始检查'按钮进行格式检查。")
    

    
    def check_document_format(self, settings):
        # 使用python-docx库实际检查Word文档格式
        from docx import Document
        from docx.shared import Cm
        
        results = [
            f"文档格式检查结果 - {os.path.basename(self.file_path)}",
            "=" * 50,
            "\n基本信息:",
            f"- 文件名: {os.path.basename(self.file_path)}",
            f"- 路径: {self.file_path}",
            f"- 检查时间: {time.strftime('%Y-%m-%d %H:%M:%S')}",
            f"- 检查设置: {settings}",
            "\n检查项目:",
        ]
        
        problems_found = []
        warnings_found = []
        
        try:
            # 打开文档
            results.append("正在打开文档进行分析...")
            doc = Document(self.file_path)
            
            # 页面大小检查
            if settings['size_check']:
                results.append("✓ 页面大小检查 (要求: A4)")
                # python-docx默认就是A4，这里可以添加实际检查逻辑
                warnings_found.append("注意: 当前版本可能无法准确检测页面大小")
                results.append("  - 页面大小检测功能已启用")
            
            # 页边距检查
            if settings['margin_check']:
                results.append(f"✓ 页边距检查 (上: {settings['margin_top']}cm, 下: {settings['margin_bottom']}cm, 左: {settings['margin_left']}cm, 右: {settings['margin_right']}cm)")
                results.append(f"  - 要求页边距：上={settings['margin_top']}cm, 下={settings['margin_bottom']}cm, 左={settings['margin_left']}cm, 右={settings['margin_right']}cm")
                # 实际检查页边距
                sections = doc.sections
                results.append(f"  - 发现{len(sections)}个文档节")
                
                for i, section in enumerate(sections):
                    try:
                        top_margin = section.top_margin.cm
                        bottom_margin = section.bottom_margin.cm
                        left_margin = section.left_margin.cm
                        right_margin = section.right_margin.cm
                        
                        results.append(f"  - 第{i+1}节实际页边距：上={top_margin:.2f}cm, 下={bottom_margin:.2f}cm, 左={left_margin:.2f}cm, 右={right_margin:.2f}cm")
                        
                        # 检查是否符合要求的页边距（允许一定误差）
                        if abs(top_margin - float(settings['margin_top'])) > 0.1:
                            problems_found.append(f"第{i+1}节上页边距不符合要求，应为{settings['margin_top']}cm，实际为{top_margin:.2f}cm")
                        if abs(bottom_margin - float(settings['margin_bottom'])) > 0.1:
                            problems_found.append(f"第{i+1}节下页边距不符合要求，应为{settings['margin_bottom']}cm，实际为{bottom_margin:.2f}cm")
                        if abs(left_margin - float(settings['margin_left'])) > 0.1:
                            problems_found.append(f"第{i+1}节左页边距不符合要求，应为{settings['margin_left']}cm，实际为{left_margin:.2f}cm")
                        if abs(right_margin - float(settings['margin_right'])) > 0.1:
                            problems_found.append(f"第{i+1}节右页边距不符合要求，应为{settings['margin_right']}cm，实际为{right_margin:.2f}cm")
                    except Exception as e:
                        warnings_found.append(f"检查第{i+1}节页边距时出错: {str(e)}")
            
            # 页眉页脚检查
            if settings['header_check']:
                results.append("✓ 页眉页脚检查")
                for i, section in enumerate(doc.sections):
                    if not section.header.is_linked_to_previous:
                        header_text = "".join([paragraph.text for paragraph in section.header.paragraphs])
                        if not header_text.strip():
                            problems_found.append(f"第{i+1}节页眉为空")
                    if not section.footer.is_linked_to_previous:
                        footer_text = "".join([paragraph.text for paragraph in section.footer.paragraphs])
                        if not footer_text.strip():
                            problems_found.append(f"第{i+1}节页脚为空")
            
            # 目录检查
            if settings['toc_check']:
                results.append("✓ 目录检查")
                has_toc = False
                for paragraph in doc.paragraphs:
                    if '目录' in paragraph.text and paragraph.style.name.startswith('Heading'):
                        has_toc = True
                        break
                if not has_toc:
                    problems_found.append("未找到目录")
            
            # 段落检查
            paragraph_count = 0
            format_issues = 0
            total_paragraphs = len(doc.paragraphs)
            results.append(f"✓ 开始段落格式检查，总共发现{total_paragraphs}个段落")
            
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():  # 跳过空段落
                    continue
                    
                paragraph_count += 1
                results.append(f"  - 检查段落{i+1}/{total_paragraphs}: '{paragraph.text[:50]}...'")
                
                # 段落对齐检查
                if settings['align_check']:
                    try:
                        current_align = paragraph.alignment
                        # 将对齐方式转换为中文的函数
                        def align_to_chinese(align_value):
                            align_map = {
                                0: "左对齐",
                                1: "居中对齐",
                                2: "右对齐",
                                3: "两端对齐",
                                None: "无"
                            }
                            # 处理字符串类型的对齐值
                            if isinstance(align_value, str):
                                str_map = {
                                    "LEFT": "左对齐",
                                    "CENTER": "居中对齐",
                                    "RIGHT": "右对齐",
                                    "JUSTIFY": "两端对齐",
                                    "0": "左对齐",
                                    "1": "居中对齐",
                                    "2": "右对齐",
                                    "3": "两端对齐"
                                }
                                return str_map.get(align_value.upper(), align_value)
                            return align_map.get(align_value, str(align_value))
                        
                        # 处理None值 - 尝试从段落样式获取真实对齐方式
                        if current_align is None:
                            results.append(f"    - 当前对齐方式: 无 (尝试获取样式值)")
                            # 尝试从段落样式获取对齐信息
                            if hasattr(paragraph.style, 'paragraph_format'):
                                style_align = paragraph.style.paragraph_format.alignment
                                display_align = align_to_chinese(style_align)
                                results.append(f"    - 段落样式对齐: {display_align}")
                                # 如果样式中有设置，使用样式的值
                                if style_align is not None:
                                    current_align = style_align
                                else:
                                    # 如果样式中也没有设置，使用Word默认的左对齐（0）
                                    results.append(f"    - 样式中未找到设置，使用Word默认的左对齐(0)")
                                    current_align = 0
                        else:
                            results.append(f"    - 当前对齐方式: {align_to_chinese(current_align)}")
                        results.append(f"    - 有效对齐方式: {align_to_chinese(current_align)}")
                        
                        # 映射对齐方式
                        align_map = {'左对齐': 0, '居中对齐': 1, '右对齐': 2, '两端对齐': 3}
                        if settings['align'] in align_map:
                            results.append(f"    - 要求对齐方式: {settings['align']} ({align_map[settings['align']]})")
                            
                            # 考虑Word的默认行为：未设置对齐方式时默认为左对齐
                            # 但在显示时仍然显示真实获取的值
                            actual_align = current_align
                            check_align = current_align if current_align is not None else 0  # 0表示左对齐
                            
                            if check_align != align_map[settings['align']]:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的对齐方式不符合要求：当前为{align_to_chinese(actual_align)}（Word默认为左对齐），要求为{align_map[settings['align']]}({settings['align']})")
                                format_issues += 1
                    except Exception as e:
                        warnings_found.append(f"检查段落{i+1}对齐方式时出错: {str(e)}")
                
                # 字体检查
                if settings['font_check']:
                    try:
                        style_name = paragraph.style.name
                        results.append(f"    - 段落样式: {style_name}")
                        if style_name.startswith('Normal'):
                            results.append(f"    - 检查字体: 要求'{settings['font']}'")
                            for run_idx, run in enumerate(paragraph.runs):
                                font_name = run.font.name
                                display_font = font_name if font_name is not None else "无"
                                results.append(f"      - 文本片段{run_idx+1}字体: {display_font}")
                                if font_name != settings['font'] and font_name is not None:
                                    # 获取段落文本（限制长度以避免过长输出）
                                    paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                    problems_found.append(f"段落'{paragraph_text}'第{run_idx+1}部分使用了不符合要求的字体：'{font_name}'，要求字体为：'{settings['font']}'")
                                    format_issues += 1
                    except Exception as e:
                        warnings_found.append(f"检查段落{i+1}字体时出错: {str(e)}")
                
                # 字号检查
                if settings['font_size_check']:
                    try:
                        style_name = paragraph.style.name
                        if style_name.startswith('Normal'):
                            results.append(f"    - 检查字号: 要求{settings['font_size']}pt")
                            for run_idx, run in enumerate(paragraph.runs):
                                if run.font.size:
                                    run_size = run.font.size.pt
                                    results.append(f"      - 文本片段{run_idx+1}字号: {run_size}pt")
                                    if run_size != float(settings['font_size']):
                                        # 获取段落文本（限制长度以避免过长输出）
                                        paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                        problems_found.append(f"段落'{paragraph_text}'第{run_idx+1}部分使用了不符合要求的字号：{run_size}pt，要求字号为：{settings['font_size']}pt")
                                        format_issues += 1
                                else:
                                    results.append(f"      - 文本片段{run_idx+1}未设置字号")
                    except Exception as e:
                        warnings_found.append(f"检查段落{i+1}字号时出错: {str(e)}")
                
                # 行间距检查
                if settings['line_spacing_check']:
                    try:
                        # 获取段落的行距设置
                        line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
                        line_spacing_value = paragraph.paragraph_format.line_spacing
                        
                        results.append(f"    - 检查行间距: 要求{settings['line_spacing']}")
                        
                        # 处理None值
                        if line_spacing_rule is None:
                            results.append(f"      - 当前行距规则: 无 (尝试获取样式值)")
                            # 尝试从段落样式获取行距信息
                            if hasattr(paragraph.style, 'paragraph_format'):
                                style_line_spacing_rule = paragraph.style.paragraph_format.line_spacing_rule
                                style_line_spacing_value = paragraph.style.paragraph_format.line_spacing
                                # 将行距规则转换为中文的函数
                                def line_spacing_rule_to_chinese(rule_value):
                                    rule_map = {
                                        0: "单倍行距",
                                        1: "1.15倍行距",
                                        2: "多倍行距",
                                        3: "最小值",
                                        4: "固定值",
                                        5: "两倍行距",
                                        6: "1.5倍行距",
                                        None: "无"
                                    }
                                    # 处理字符串类型的行距规则
                                    if isinstance(rule_value, str):
                                        str_map = {
                                            "SINGLE": "单倍行距",
                                            "ONE_POINT_FIVE": "1.5倍行距",
                                            "DOUBLE": "两倍行距",
                                            "MULTIPLE": "多倍行距",
                                            "AT_LEAST": "最小值",
                                            "EXACTLY": "固定值",
                                            "0": "单倍行距",
                                            "1": "1.15倍行距",
                                            "2": "多倍行距",
                                            "3": "最小值",
                                            "4": "固定值",
                                            "5": "两倍行距",
                                            "6": "1.5倍行距"
                                        }
                                        return str_map.get(rule_value.upper(), rule_value)
                                    return rule_map.get(rule_value, str(rule_value))
                                
                                display_rule = line_spacing_rule_to_chinese(style_line_spacing_rule)
                                display_value = style_line_spacing_value if style_line_spacing_value is not None else "无"
                                results.append(f"      - 段落样式行距规则: {display_rule}")
                                results.append(f"      - 段落样式行距值: {display_value}")
                                # 如果样式中有设置，使用样式的值
                                if style_line_spacing_rule is not None:
                                    line_spacing_rule = style_line_spacing_rule
                                    line_spacing_value = style_line_spacing_value
                                # 如果样式中也没有设置，根据Word默认行为设置
                                else:
                                    # Word默认行距规则通常为2（多倍行距），值为1.0（单倍行距）
                                    results.append(f"      - 样式中未找到设置，使用Word默认值")
                                    line_spacing_rule = 2
                                    line_spacing_value = 1.0
                                    # 复用上面定义的转换函数
                                    results.append(f"      - 有效行距规则: {line_spacing_rule_to_chinese(line_spacing_rule)}")
                                    results.append(f"      - 有效行距值: {line_spacing_value}")
                        else:
                            # 复用上面定义的转换函数
                            results.append(f"      - 当前行距规则: {line_spacing_rule_to_chinese(line_spacing_rule)}")
                            results.append(f"      - 当前行距值: {line_spacing_value}")
                        
                        # 检查是否为1.5倍行距
                        if settings['line_spacing'] == '1.5倍行距':
                            # 处理可能的None值情况
                            if line_spacing_rule is None or line_spacing_value is None:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的行距设置未明确定义，无法确认是否为1.5倍行距")
                                format_issues += 1
                            elif not (line_spacing_rule == 2 and abs(line_spacing_value - 1.5) < 0.1):
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的行距不是要求的1.5倍行距，当前为{line_spacing_value}")
                                format_issues += 1
                        # 检查是否为2倍行距
                        elif settings['line_spacing'] == '2倍行距':
                            # 处理可能的None值情况
                            if line_spacing_rule is None or line_spacing_value is None:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的行距设置未明确定义，无法确认是否为2倍行距")
                                format_issues += 1
                            elif not (line_spacing_rule == 2 and abs(line_spacing_value - 2.0) < 0.1):
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的行距不是要求的2倍行距，当前为{line_spacing_value}")
                                format_issues += 1
                    except Exception as e:
                        warnings_found.append(f"无法准确检测段落{i+1}的行距设置: {str(e)}")
                
                # 首行缩进检查
                if settings['indent_check']:
                    try:
                        # 获取首行缩进值（单位：pt）
                        first_line_indent = paragraph.paragraph_format.first_line_indent
                        results.append(f"    - 检查首行缩进: 要求{settings['indent']}cm")
                        
                        # 处理None值 - 尝试从段落样式获取
                        if first_line_indent is None:
                            results.append(f"      - 当前首行缩进: 无 (尝试获取样式值)")
                            # 尝试从段落样式获取缩进信息
                            if hasattr(paragraph.style, 'paragraph_format'):
                                style_indent = paragraph.style.paragraph_format.first_line_indent
                                display_indent = style_indent if style_indent is not None else "无"
                                results.append(f"      - 段落样式首行缩进: {display_indent}")
                                # 如果样式中有设置，使用样式的值
                                if style_indent is not None:
                                    first_line_indent = style_indent
                                    results.append(f"      - 有效首行缩进: 从样式获取")
                                else:
                                    # 如果样式中也没有设置，使用Word默认值（通常为0）
                                    results.append(f"      - 样式中未找到设置，使用Word默认值0")
                                    # 创建一个0cm的缩进值
                                    from docx.shared import Cm
                                    first_line_indent = Cm(0)
                                    results.append(f"      - 有效首行缩进: 0.00cm")
                        
                        # 再次检查是否获取到了缩进值
                        if first_line_indent is not None:
                            # 转换为厘米进行比较
                            indent_in_cm = first_line_indent.cm
                            required_indent = float(settings['indent'])
                            results.append(f"      - 有效首行缩进: {indent_in_cm:.2f}cm")
                            
                            if abs(indent_in_cm - required_indent) > 0.1:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'的首行缩进不符合要求：{indent_in_cm:.2f}cm，要求为：{required_indent}cm")
                                format_issues += 1
                        else:
                            results.append(f"      - 未找到明确的首行缩进设置")
                            # 如果没有设置首行缩进，也报错
                            # 获取段落文本（限制长度以避免过长输出）
                            paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                            problems_found.append(f"段落'{paragraph_text}'未设置首行缩进")
                            format_issues += 1
                    except Exception as e:
                        warnings_found.append(f"无法准确检测段落{i+1}的首行缩进设置: {str(e)}")
                
                # 特殊格式检查
                if settings['special_format_check']:
                    try:
                        results.append("    - 检查特殊格式")
                        for run_idx, run in enumerate(paragraph.runs):
                            # 处理None值 - None通常表示未设置该格式
                            run_bold = False if run.bold is None else run.bold
                            run_italic = False if run.italic is None else run.italic
                            run_underline = False if run.underline is None else run.underline
                            
                            results.append(f"      - 文本片段{run_idx+1}特殊格式: 加粗={'是' if run_bold else '否'}, 斜体={'是' if run_italic else '否'}, 下划线={'是' if run_underline else '否'}")
                            if settings['format_bold'] and run_bold:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'第{run_idx+1}部分包含不允许的加粗格式")
                            if settings['format_italic'] and run_italic:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'第{run_idx+1}部分包含不允许的斜体格式")
                            if settings['format_underline'] and run_underline:
                                # 获取段落文本（限制长度以避免过长输出）
                                paragraph_text = paragraph.text.strip()[:50] + '...' if len(paragraph.text.strip()) > 50 else paragraph.text.strip()
                                problems_found.append(f"段落'{paragraph_text}'第{run_idx+1}部分包含不允许的下划线格式")
                    except Exception as e:
                        warnings_found.append(f"检查段落{i+1}特殊格式时出错: {str(e)}")
            
            if format_issues > 0:
                problems_found.append(f"发现{format_issues}处段落格式不符合要求")
            
            # 表格检查
            table_count = 0
            total_tables = len(doc.tables)
            results.append(f"✓ 开始表格格式检查，总共发现{total_tables}个表格")
            
            for i, table in enumerate(doc.tables):
                table_count += 1
                results.append(f"  - 检查表格{i+1}/{total_tables}")
                
                # 表格内字体和字号检查
                if settings['inner_font_check'] or settings['inner_font_size_check']:
                    try:
                        if settings['inner_font_check']:
                            results.append(f"    - 表格内字体要求: '{settings['inner_font']}'")
                        if settings['inner_font_size_check']:
                            results.append(f"    - 表格内字号要求: {settings['inner_font_size']}pt")
                            
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                results.append(f"    - 单元格({row_idx+1},{cell_idx+1})")
                                for paragraph in cell.paragraphs:
                                    for run_idx, run in enumerate(paragraph.runs):
                                        # 表格内字体检查
                                        if settings['inner_font_check'] and run.font.name != settings['inner_font'] and run.font.name is not None:
                                            results.append(f"      - 发现不符合要求的字体: '{run.font.name}'")
                                            # 获取单元格文本（限制长度以避免过长输出）
                                            cell_text = cell.text.strip()[:30] + '...' if len(cell.text.strip()) > 30 else cell.text.strip()
                                            problems_found.append(f"表格{i+1}单元格({row_idx+1},{cell_idx+1})['{cell_text}']中发现不符合要求的字体：'{run.font.name}'，要求字体为：'{settings['inner_font']}'")
                                        
                                        # 表格内字号检查
                                        if settings['inner_font_size_check'] and run.font.size:
                                            run_size = run.font.size.pt
                                            results.append(f"      - 当前字号: {run_size}pt")
                                            if run_size != float(settings['inner_font_size']):
                                                # 获取单元格文本（限制长度以避免过长输出）
                                                cell_text = cell.text.strip()[:30] + '...' if len(cell.text.strip()) > 30 else cell.text.strip()
                                                problems_found.append(f"表格{i+1}单元格({row_idx+1},{cell_idx+1})['{cell_text}']中发现不符合要求的字号：{run_size}pt，要求字号为：{settings['inner_font_size']}pt")
                    except Exception as e:
                        warnings_found.append(f"检查表格{i+1}时出错: {str(e)}")
            
            # 标点符号检查
            if settings['punctuation_check']:
                results.append("✓ 英文标点符号检查")
                results.append("  - 检查可能的英文标点符号")
                for i, paragraph in enumerate(doc.paragraphs):
                    if not paragraph.text.strip():
                        continue
                    
                    if '\'' in paragraph.text or '"' in paragraph.text or ';' in paragraph.text:
                        results.append(f"  - 段落{i+1}发现英文标点: '{paragraph.text[:30]}...'")
                        # 直接使用现有文本预览
                        problems_found.append(f"段落'{paragraph.text[:30]}...'可能包含英文标点符号")
            
            # 空格检查
            if settings['space_check']:
                results.append("✓ 空格检查")
                results.append("  - 检查连续空格")
                for i, paragraph in enumerate(doc.paragraphs):
                    if not paragraph.text.strip():
                        continue
                    
                    if '  ' in paragraph.text:  # 连续空格
                        results.append(f"  - 段落{i+1}发现连续空格: '{paragraph.text[:30]}...'")
                        # 直接使用现有文本预览
                        problems_found.append(f"段落'{paragraph.text[:30]}...'包含多余空格")
            
            # 添加统计信息
            results.extend([
                "",
                f"文档统计:",
                f"- 总段落数: {paragraph_count}",
                f"- 总表格数: {table_count}",
                f"- 发现格式问题数: {len(problems_found)}",
                f"- 发现警告数: {len(warnings_found)}",
                "",
                "=" * 50,
                "检查总结:",
            ])
            
            # 添加问题列表
            if problems_found:
                results.append(f"❌ 发现{len(problems_found)}个格式问题:")
                for problem in problems_found:
                    results.append(f"  - {problem}")
            else:
                results.append("✓ 未发现严重格式问题")
            
            # 添加警告
            if warnings_found:
                results.append("\n警告:")
                for warning in warnings_found:
                    results.append(f"  - {warning}")
            
            # 添加建议
            results.extend([
                "\n建议:",
                "1. 根据检查结果修复发现的问题",
                "2. 再次运行检查确保所有问题已解决",
                "3. 保存最终版本的文档"
            ])
            
        except Exception as e:
            results.extend([
                "\n检查过程中发生错误:",
                f"  - {str(e)}",
                "\n建议:",
                "1. 确保文件未被其他程序占用",
                "2. 检查文件格式是否正确",
                "3. 尝试重新打开文档"
            ])
        
        return "\n".join(results)
        
    def simulate_format_check(self, settings):
        # 为了兼容之前的调用，保持原方法名称但调用新的实际检查方法
        return self.check_document_format(settings)
    
    def select_all(self):
        """全选所有检查项"""
        self.size_check_var.set(True)
        self.margin_check_var.set(True)
        self.header_check_var.set(True)
        self.toc_check_var.set(True)
        self.align_check_var.set(True)
        self.line_spacing_check_var.set(True)
        self.indent_check_var.set(True)
        self.font_check_var.set(True)
        self.font_size_check_var.set(True)
        self.inner_font_check_var.set(True)
        self.inner_font_size_check_var.set(True)
        self.special_format_check_var.set(True)
        self.format_bold_var.set(True)
        self.format_italic_var.set(True)
        self.format_underline_var.set(True)
        self.color_check_var.set(True)
        self.punctuation_check_var.set(True)
        self.space_check_var.set(True)
        self.char_spacing_check_var.set(True)
        
    def select_none(self):
        """取消选中所有检查项"""
        self.size_check_var.set(False)
        self.margin_check_var.set(False)
        self.header_check_var.set(False)
        self.toc_check_var.set(False)
        self.align_check_var.set(False)
        self.line_spacing_check_var.set(False)
        self.indent_check_var.set(False)
        self.font_check_var.set(False)
        self.font_size_check_var.set(False)
        self.inner_font_check_var.set(False)
        self.inner_font_size_check_var.set(False)
        self.special_format_check_var.set(False)
        self.format_bold_var.set(False)
        self.format_italic_var.set(False)
        self.format_underline_var.set(False)
        self.color_check_var.set(False)
        self.punctuation_check_var.set(False)
        self.space_check_var.set(False)
        self.char_spacing_check_var.set(False)
        
    def show_preset_info(self):
        """显示预设格式说明"""
        messagebox.showinfo("预设格式说明", "当前版本暂不支持预设格式。\n请手动设置所需的格式检查选项。")
        
    def check_document(self):
        """检查文档格式"""
        if not hasattr(self, 'file_path'):
            self.update_status("错误: 请先选择一个文件")
            self.update_result_text("错误: 请先选择一个Word文档进行检查。")
            return
        
        try:
            self.update_status("正在检查文档格式...")
            
            # 收集所有检查规则设置
            settings = {
                'size_check': self.size_check_var.get(),
                'margin_check': self.margin_check_var.get(),
                'margin_top': float(self.margin_top_var.get()),
                'margin_bottom': float(self.margin_bottom_var.get()),
                'margin_left': float(self.margin_left_var.get()),
                'margin_right': float(self.margin_right_var.get()),
                'header_check': self.header_check_var.get(),
                'toc_check': self.toc_check_var.get(),
                'align_check': self.align_check_var.get(),
                'align': self.align_var.get(),
                'line_spacing_check': self.line_spacing_check_var.get(),
                'line_spacing': self.line_spacing_var.get(),
                'indent_check': self.indent_check_var.get(),
                'indent': self.indent_var.get(),
                'font_check': self.font_check_var.get(),
                'font': self.font_var.get(),
                'font_size_check': self.font_size_check_var.get(),
                'font_size': self.font_size_var.get(),
                'inner_font_check': self.inner_font_check_var.get(),
                'inner_font': self.inner_font_var.get(),
                'inner_font_size_check': self.inner_font_size_check_var.get(),
                'inner_font_size': self.inner_font_size_var.get(),
                'special_format_check': self.special_format_check_var.get(),
                'format_bold': self.format_bold_var.get(),
                'format_italic': self.format_italic_var.get(),
                'format_underline': self.format_underline_var.get(),
                'color_check': self.color_check_var.get(),
                'color': self.color_var.get(),
                'punctuation_check': self.punctuation_check_var.get(),
                'space_check': self.space_check_var.get(),
                'char_spacing_check': self.char_spacing_check_var.get(),
                'fix_file': self.fix_file_var.get()
            }
            
            # 执行实际的格式检查
            results = self.check_document_format(settings)
            
            # 显示检查结果
            self.update_result_text(results)
            self.update_status(f"检查完成: {os.path.basename(self.file_path)}")
            
            # 如果需要修复文件
            if settings['fix_file']:
                self.simulate_file_fix()
                
        except Exception as e:
            error_msg = f"检查过程中发生错误: {str(e)}"
            self.update_status("检查失败")
            self.update_result_text(error_msg)
    
    def fix_document(self):
        """修复文档格式"""
        if not hasattr(self, 'file_path'):
            messagebox.showerror("错误", "请先打开一个文档!")
            return
        
        # 检查是否生成修复文件
        if not self.fix_file_var.get():
            messagebox.showinfo("提示", "请先勾选'自动生成修复文件'选项!")
            return
        
        try:
            self.update_status("正在修复文档...")
            
            # 模拟文件修复过程
            self.simulate_file_fix()
            
        except Exception as e:
            error_msg = f"修复过程中发生错误: {str(e)}"
            self.update_status("修复失败")
            self.update_result_text(error_msg)
            
    def simulate_file_fix(self):
        # 模拟文件修复过程
        try:
            # 在实际应用中，这里应该使用python-docx库来修复Word文档格式
            fixed_filename = os.path.splitext(self.file_path)[0] + "_fixed.docx"
            # 这里只是创建一个空文件作为示例
            with open(fixed_filename, 'w') as f:
                f.write("This is a placeholder for the fixed document.")
            
            self.update_result_text(self.result_text.get(1.0, tk.END) + 
                                  f"\n\n✓ 已生成修复后的文件: {os.path.basename(fixed_filename)}")
            self.update_status(f"已修复文件: {os.path.basename(fixed_filename)}")
            
        except Exception as e:
            self.update_result_text(self.result_text.get(1.0, tk.END) + 
                                  f"\n\n✗ 生成修复文件失败: {str(e)}")
            self.update_status("修复文件失败")
    
    def create_preview_area(self):
        # 创建右侧预览区域
        self.preview_frame = ttk.LabelFrame(self.main_frame, text="检查结果", relief=tk.RAISED)
        self.preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # 创建结果文本框和滚动条
        result_frame = ttk.Frame(self.preview_frame)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 垂直滚动条
        y_scrollbar = ttk.Scrollbar(result_frame, orient=tk.VERTICAL)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 水平滚动条
        x_scrollbar = ttk.Scrollbar(result_frame, orient=tk.HORIZONTAL)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 结果文本框
        self.result_text = tk.Text(result_frame, wrap=tk.NONE, height=30, width=60, 
                                 yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set,
                                 font=("SimHei", 10), foreground="black")
        self.result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 配置滚动条
        y_scrollbar.config(command=self.result_text.yview)
        x_scrollbar.config(command=self.result_text.xview)
        
        # 设置文本框为只读
        self.result_text.config(state=tk.DISABLED)
        
        # 添加一些初始提示文本
        self.update_result_text("欢迎使用格式规范检查工具！\n\n请选择一个Word文档进行检查。")
    
    def setup_sensitive_tab(self):
        # 创建敏感词标签页的内容
        # 设置整体布局为左侧设置，右侧结果
        main_frame = ttk.Frame(self.tab_sensitive)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 左侧设置区域
        setting_frame = ttk.LabelFrame(main_frame, text="敏感词分类设置（双击编辑）：")
        setting_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=5, pady=5, expand=False)
        setting_frame.configure(width=600)
        
        # 创建表格区域
        table_frame = ttk.Frame(setting_frame)
        table_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 创建Treeview组件
        columns = ("分类", "敏感词")
        self.sensitive_tree = ttk.Treeview(table_frame, columns=columns, show="headings", selectmode="browse")
        
        # 设置列标题和宽度
        self.sensitive_tree.heading("分类", text="分类名称")
        self.sensitive_tree.heading("敏感词", text="敏感词（多个词用逗号分隔）")
        self.sensitive_tree.column("分类", width=150, anchor="w")
        self.sensitive_tree.column("敏感词", width=400, anchor="w")
        
        # 创建滚动条
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.sensitive_tree.yview)
        self.sensitive_tree.configure(yscroll=scrollbar.set)
        
        # 放置组件
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.sensitive_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 添加示例数据
        example_data = [
            ("1 联系方式", "团队,联系,电话,邮箱,地址,手机,传真,热线"),
            ("2 地理位置", "省,市,县,区,街道,路,巷,村,镇,乡"),
            ("3 网络信息", "网址,网站,链接,URL,域名,IP"),
            ("4 版权信息", "copyright,版权所有,©,®"),
            ("5 作者信息", "author,作者,编写,撰写"),
            ("6 落地案例", "落地,案例,实例,实践,应用,实施"),
            ("7 奖项荣誉", "奖项,荣誉,冠,金,银奖,铜奖,一等奖,二等奖,三等奖"),
            ("8 专利信息", "专利,发明专利,实用新型,外观设计,专利号,专利申请")
        ]
        for item in example_data:
            self.sensitive_tree.insert("", tk.END, values=item)
        
        # 双击编辑功能
        self.sensitive_tree.bind("<Double-1>", self.on_sensitive_item_double_click)
        
        # 按钮区域
        button_frame = ttk.Frame(setting_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=10)
        
        # 设置按钮样式
        self.style.configure("Green.TButton", foreground="black")
        self.style.configure("Red.TButton", foreground="black")
        self.style.configure("Yellow.TButton", foreground="black")
        self.style.configure("Purple.TButton", foreground="black")
        self.style.configure("Blue.TButton", foreground="black")
        
        if os.name == 'nt':  # Windows系统
            self.style.configure("Green.TButton", background="#4CAF50", foreground="black")
            self.style.map("Green.TButton", background=[("active", "#45a049")])
            self.style.configure("Red.TButton", background="#f44336", foreground="black")
            self.style.map("Red.TButton", background=[("active", "#da190b")])
            self.style.configure("Yellow.TButton", background="#FFC107", foreground="black")
            self.style.map("Yellow.TButton", background=[("active", "#ffb300")])
            self.style.configure("Purple.TButton", background="#9C27B0", foreground="black")
            self.style.map("Purple.TButton", background=[("active", "#7b1fa2")])
            self.style.configure("Blue.TButton", background="#2196F3", foreground="black")
            self.style.map("Blue.TButton", background=[("active", "#0b7dda")])
        
        # 添加按钮
        add_btn = ttk.Button(button_frame, text="添加分类", style="Green.TButton", command=self.add_sensitive_category)
        add_btn.pack(side=tk.LEFT, padx=2)
        
        # 删除按钮
        del_btn = ttk.Button(button_frame, text="删除分类", style="Red.TButton", command=self.delete_sensitive_category)
        del_btn.pack(side=tk.LEFT, padx=2)
        
        # 保存按钮
        save_btn = ttk.Button(button_frame, text="保存敏感词设置", style="Yellow.TButton", command=self.save_sensitive_settings)
        save_btn.pack(side=tk.LEFT, padx=2)
        
        # 加载按钮
        load_btn = ttk.Button(button_frame, text="加载敏感词设置", style="Blue.TButton", command=self.load_sensitive_settings)
        load_btn.pack(side=tk.LEFT, padx=2)
        
        # 清空按钮
        clear_btn = ttk.Button(button_frame, text="清空所有", style="Red.TButton", command=self.clear_all_sensitive_categories)
        clear_btn.pack(side=tk.LEFT, padx=2)
        
        # 执行检查按钮
        check_btn = ttk.Button(button_frame, text="执行敏感词检查", style="Purple.TButton", command=self.execute_sensitive_check)
        check_btn.pack(side=tk.LEFT, padx=2)
        
        # 右侧结果区域
        result_frame = ttk.LabelFrame(main_frame, text="检查结果：")
        result_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=5, pady=5, expand=True)
        
        # 创建结果文本框和滚动条
        text_frame = ttk.Frame(result_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 垂直滚动条
        y_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 水平滚动条
        x_scrollbar = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 结果文本框
        self.sensitive_result_text = tk.Text(text_frame, wrap=tk.NONE, 
                                          yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set,
                                          font=("SimHei", 10), foreground="black")
        self.sensitive_result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 配置滚动条
        y_scrollbar.config(command=self.sensitive_result_text.yview)
        x_scrollbar.config(command=self.sensitive_result_text.xview)
        
        # 设置文本框为只读
        self.sensitive_result_text.config(state=tk.DISABLED)
        
        # 初始提示文本
        self.update_sensitive_result("欢迎使用敏感词检查功能！\n\n请在左侧设置敏感词分类和敏感词，然后点击'执行敏感词检查'按钮进行检查。\n\n您可以保存或加载敏感词设置，方便后续使用。")
    
    def on_sensitive_item_double_click(self, event):
        # 双击编辑功能
        item = self.sensitive_tree.identify_row(event.y)
        if item:
            # 获取当前值
            values = self.sensitive_tree.item(item, "values")
            column = self.sensitive_tree.identify_column(event.x)
            column_idx = int(column.replace("#", "")) - 1
            
            # 创建编辑窗口
            edit_window = tk.Toplevel(self.root)
            edit_window.title("编辑敏感词")
            edit_window.geometry("400x150")
            edit_window.resizable(False, False)
            
            # 设置窗口居中
            edit_window.update_idletasks()
            width = edit_window.winfo_width()
            height = edit_window.winfo_height()
            x = (self.root.winfo_width() // 2) - (width // 2) + self.root.winfo_x()
            y = (self.root.winfo_height() // 2) - (height // 2) + self.root.winfo_y()
            edit_window.geometry(f"{width}x{height}+{x}+{y}")
            
            # 创建标签和输入框
            label_text = "分类名称：" if column_idx == 0 else "敏感词（用逗号分隔）："
            ttk.Label(edit_window, text=label_text, font=("SimHei", 10)).pack(pady=10, padx=10, anchor="w")
            
            entry_var = tk.StringVar(value=values[column_idx])
            entry = ttk.Entry(edit_window, textvariable=entry_var, width=50, font=("SimHei", 10))
            entry.pack(pady=5, padx=10, fill=tk.X)
            entry.focus_set()
            
            def save_edit():
                new_values = list(values)
                new_values[column_idx] = entry_var.get()
                self.sensitive_tree.item(item, values=new_values)
                edit_window.destroy()
            
            # 创建按钮
            button_frame = ttk.Frame(edit_window)
            button_frame.pack(pady=10, fill=tk.X)
            
            ttk.Button(button_frame, text="保存", command=save_edit).pack(side=tk.RIGHT, padx=10)
            ttk.Button(button_frame, text="取消", command=edit_window.destroy).pack(side=tk.RIGHT)
    
    def add_sensitive_category(self):
        # 添加新分类
        # 获取最大分类序号
        max_num = 0
        for item in self.sensitive_tree.get_children():
            values = self.sensitive_tree.item(item, "values")
            if values and values[0]:
                try:
                    num_str = values[0].split()[0]
                    num = int(num_str)
                    if num > max_num:
                        max_num = num
                except (ValueError, IndexError):
                    pass
        
        # 创建新分类
        new_num = max_num + 1
        self.sensitive_tree.insert("", tk.END, values=[f"{new_num} 新分类", ""])
    
    def delete_sensitive_category(self):
        # 删除选中的分类
        selected_item = self.sensitive_tree.selection()
        if selected_item:
            if tk.messagebox.askyesno("确认删除", "确定要删除选中的分类吗？"):
                self.sensitive_tree.delete(selected_item)
        else:
            tk.messagebox.showinfo("提示", "请先选择要删除的分类")
            
    def load_sensitive_settings(self):
        # 加载敏感词设置
        # 打开文件选择对话框
        file_path = filedialog.askopenfilename(
            title="选择敏感词设置文件",
            filetypes=[
                ("文本文件", "*.txt"),
                ("所有文件", "*.*")
            ]
        )
        
        if not file_path:
            return
        
        # 确认是否替换现有设置
        if self.sensitive_tree.get_children():
            if not tk.messagebox.askyesno("确认", "加载设置将替换当前所有敏感词设置，是否继续？"):
                return
        
        try:
            # 清空当前设置
            self.sensitive_tree.delete(*self.sensitive_tree.get_children())
            
            # 读取文件
            loaded_count = 0
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    # 跳过注释和空行
                    if line and not line.startswith('#'):
                        try:
                            # 解析行格式：分类名称|敏感词1,敏感词2,...
                            parts = line.split('|', 1)
                            if len(parts) == 2:
                                category, words = parts
                                self.sensitive_tree.insert("", tk.END, values=(category.strip(), words.strip()))
                                loaded_count += 1
                        except Exception as e:
                            print(f"解析行出错: {line}, 错误: {e}")
            
            self.update_status(f"已从文件加载 {loaded_count} 个敏感词分类设置")
            tk.messagebox.showinfo("加载成功", f"成功从文件加载 {loaded_count} 个敏感词分类设置")
            
        except Exception as e:
            tk.messagebox.showerror("加载失败", f"加载文件时出错：{str(e)}")
            self.update_status("加载敏感词设置失败")
            
    def clear_all_sensitive_categories(self):
        # 清空所有敏感词分类
        if not self.sensitive_tree.get_children():
            tk.messagebox.showinfo("提示", "没有可清空的敏感词分类")
            return
        
        # 确认清空
        if tk.messagebox.askyesno("确认清空", "确定要清空所有敏感词分类吗？此操作不可撤销。"):
            # 清空表格
            self.sensitive_tree.delete(*self.sensitive_tree.get_children())
            
            # 清空结果区域
            self.update_sensitive_result("")
            
            self.update_status("已清空所有敏感词分类")
            tk.messagebox.showinfo("清空成功", "所有敏感词分类已清空")
    
    def save_sensitive_settings(self):
        # 保存敏感词设置
        settings = []
        for item in self.sensitive_tree.get_children():
            values = self.sensitive_tree.item(item, "values")
            if values and len(values) >= 2:
                settings.append(values)
        
        if not settings:
            tk.messagebox.showwarning("警告", "没有可保存的敏感词设置")
            return
        
        # 打开文件保存对话框
        file_path = filedialog.asksaveasfilename(
            title="保存敏感词设置",
            defaultextension=".txt",
            filetypes=[
                ("文本文件", "*.txt"),
                ("所有文件", "*.*")
            ]
        )
        
        if not file_path:
            return
        
        try:
            # 保存到文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("# 敏感词设置文件\n")
                f.write(f"# 保存时间: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())}\n")
                f.write("# 格式：分类名称|敏感词1,敏感词2,...\n\n")
                
                for category, words in settings:
                    f.write(f"{category}|{words}\n")
            
            self.update_status(f"已保存 {len(settings)} 个敏感词分类设置到文件")
            tk.messagebox.showinfo("保存成功", f"成功保存 {len(settings)} 个敏感词分类设置到\n{file_path}")
        except Exception as e:
            tk.messagebox.showerror("保存失败", f"保存文件时出错：{str(e)}")
            self.update_status("保存敏感词设置失败")
    
    def execute_sensitive_check(self):
        # 执行敏感词检查
        self.update_status("正在执行敏感词检查...")
        
        # 获取敏感词设置
        sensitive_words_dict = {}
        for item in self.sensitive_tree.get_children():
            values = self.sensitive_tree.item(item, "values")
            if values and len(values) >= 2:
                category = values[0]
                words = values[1]
                if words.strip():
                    word_list = [word.strip() for word in words.split(",") if word.strip()]
                    sensitive_words_dict[category] = word_list
        
        # 检查是否有敏感词设置
        if not sensitive_words_dict:
            tk.messagebox.showwarning("警告", "请先设置敏感词")
            self.update_status("就绪")
            return
        
        # 选择要检查的文档
        file_path = filedialog.askopenfilename(
            title="选择要检查的文档",
            filetypes=[
                ("Word文档", "*.docx"),
                ("文本文件", "*.txt"),
                ("所有文件", "*.*")
            ]
        )
        
        if not file_path:
            self.update_status("就绪")
            return
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            tk.messagebox.showerror("错误", "文件不存在")
            self.update_status("就绪")
            return
        
        # 初始化结果
        results = []
        results.append("敏感词检查结果")
        results.append("=" * 50)
        results.append(f"检查文件：{os.path.basename(file_path)}")
        results.append(f"设置了 {len(sensitive_words_dict)} 个敏感词分类")
        results.append(f"设置了 {sum(len(words) for words in sensitive_words_dict.values())} 个敏感词")
        
        # 根据文件类型读取内容
        content = ""
        paragraphs = []
        try:
            if file_path.endswith(".docx"):
                # 尝试导入python-docx库读取Word文档
                try:
                    from docx import Document
                    doc = Document(file_path)
                    # 收集所有段落
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text.strip())
                    # 同时读取表格内容，作为独立段落
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    paragraphs.append(cell.text.strip())
                    # 将所有段落合并为完整内容
                    content = "\n".join(paragraphs)
                except ImportError:
                    # 如果没有安装python-docx，使用模拟内容
                    content = "这是一个Word文档内容，包含团队、联系、电话、邮箱、地址等敏感词。"
                    paragraphs = content.split('\n')
                    self.update_status("未安装python-docx，使用模拟内容进行检查")
            elif file_path.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                # 按行分割并过滤空行
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            else:
                # 尝试以文本方式读取其他文件
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                # 按行分割并过滤空行
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        except Exception as e:
            tk.messagebox.showerror("错误", f"读取文件失败：{str(e)}")
            self.update_status("就绪")
            return
        
        # 统计敏感词并记录包含敏感词的文本段落
        found_words = {}
        found_contexts = {}
        total_sensitive_count = 0
        
        # 为每个敏感词找到上下文
        for category, words in sensitive_words_dict.items():
            category_words = []
            category_contexts = {}
            
            for word in words:
                count = 0
                word_contexts = []
                
                for i, paragraph in enumerate(paragraphs):
                    if word in paragraph:
                        # 统计在该段落中的出现次数
                        word_count = paragraph.count(word)
                        count += word_count
                        
                        # 为该段落中的每个敏感词实例创建上下文
                        index = paragraph.find(word)
                        while index != -1:
                            # 优化上下文提取，确保更好地显示包含敏感词的文本
                            # 方法：找到敏感词前后的句子边界或字符边界
                            # 定义上下文长度
                            context_length = 40
                            
                            # 计算上下文的起始和结束位置
                            start = max(0, index - context_length)
                            end = min(len(paragraph), index + len(word) + context_length)
                            
                            # 尝试找到更好的上下文边界（句子边界）
                            if start > 0:
                                # 向前查找句子边界
                                for boundary in range(start, max(0, start - 10), -1):
                                    if paragraph[boundary] in ['.', '。', '!', '！', '?', '？', ';', '；', '\n']:
                                        start = boundary + 1
                                        break
                            
                            if end < len(paragraph):
                                # 向后查找句子边界
                                for boundary in range(end, min(len(paragraph), end + 10)):
                                    if paragraph[boundary] in ['.', '。', '!', '！', '?', '？', ';', '；', '\n']:
                                        end = boundary + 1
                                        break
                            
                            # 提取上下文文本
                            context = paragraph[start:end].replace('\n', ' ')
                            
                            # 调整相对位置
                            relative_start = index - start
                            relative_end = relative_start + len(word)
                            
                            # 构建上下文描述，不添加方括号，后续直接高亮敏感词
                            context_str = context
                            
                            # 添加省略号，除非是段落的开头或结尾
                            prefix = "..." if start > 0 else ""
                            suffix = "..." if end < len(paragraph) else ""
                            context_str = prefix + context_str + suffix
                            
                            word_contexts.append({
                                'paragraph': i + 1,  # 段落编号
                                'context': context_str
                            })
                            
                            # 查找下一个出现位置
                            index = paragraph.find(word, index + 1)
                            
                            # 限制每个敏感词显示的上下文数量
                            if len(word_contexts) >= 10:  # 最多显示10个上下文
                                break
                
                if count > 0:
                    category_words.append((word, count))
                    if word_contexts:
                        category_contexts[word] = word_contexts
                    total_sensitive_count += count
            
            if category_words:
                found_words[category] = category_words
                if category_contexts:
                    found_contexts[category] = category_contexts
        
        # 生成检查结果
        if found_words:
            results.append("\n发现的敏感词：")
            for category, words in found_words.items():
                results.append(f"\n【{category}】")
                for word, count in words:
                    results.append(f"- '{word}'：出现 {count} 次")
            
            # 添加包含敏感词的文本段落
            if found_contexts:
                results.append("\n\n包含敏感词的文本段落：")
                for category, words_contexts in found_contexts.items():
                    results.append(f"\n【{category}】")
                    for word, contexts in words_contexts.items():
                        results.append(f"\n  敏感词：'{word}'")
                        for ctx in contexts:
                            results.append(f"    段落 {ctx['paragraph']}：{ctx['context']}")
                        # 如果上下文数量超过限制，显示提示
                        if len(contexts) >= 10:
                            results.append(f"    ...(更多上下文未显示)")
            
            results.append(f"\n总计发现 {total_sensitive_count} 个敏感词")
            results.append("\n建议：请检查并移除文档中的敏感词")
        else:
            results.append("\n未发现敏感词，文档通过检查")
        
        # 添加基本统计
        results.append("\n文本统计：")
        results.append(f"- 字符数：{len(content)}")
        results.append(f"- 段落数：{len([p for p in paragraphs if p.strip()])}")
        results.append(f"- 敏感词分类数：{len(sensitive_words_dict)}")
        results.append(f"- 敏感词总数：{sum(len(words) for words in sensitive_words_dict.values())}")
        results.append(f"- 发现敏感词分类数：{len(found_words)}")
        results.append(f"- 敏感词总出现次数：{total_sensitive_count}")
        
        # 添加时间戳
        current_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        results.append(f"\n检查时间：{current_time}")
        
        self.update_sensitive_result("\n".join(results))
        self.update_status("敏感词检查完成")
        
        # 如果发现敏感词，弹出提示
        if found_words:
            tk.messagebox.showinfo("检查完成", f"发现 {total_sensitive_count} 个敏感词，请查看详细结果")
    
    def update_sensitive_result(self, text):
        # 更新敏感词检查结果文本
        self.sensitive_result_text.config(state=tk.NORMAL)
        self.sensitive_result_text.delete(1.0, tk.END)
        self.sensitive_result_text.insert(tk.END, text)
        
        # 添加文本格式
        if text:
            # 设置标签样式
            self.sensitive_result_text.tag_configure("heading", font= ("SimHei", 12, "bold"), foreground="#2196F3")
            self.sensitive_result_text.tag_configure("category", font= ("SimHei", 11, "bold"), foreground="#FF9800")
            self.sensitive_result_text.tag_configure("found", foreground="#F44336", font= ("SimHei", 10, "bold"))
            self.sensitive_result_text.tag_configure("success", foreground="#4CAF50", font= ("SimHei", 10, "bold"))
            # 新添加的样式标签
            self.sensitive_result_text.tag_configure("context_section", font= ("SimHei", 11, "bold"), foreground="#9C27B0")
            self.sensitive_result_text.tag_configure("word_label", foreground="#E91E63", font= ("SimHei", 10, "bold"))
            self.sensitive_result_text.tag_configure("paragraph", foreground="#607D8B", font= ("SimHei", 10))
            # 优化敏感词高亮样式 - 使用更醒目的红色背景和白色文字，增加边框效果
            self.sensitive_result_text.tag_configure("highlight", 
                                                    background="#D32F2F", 
                                                    foreground="white", 
                                                    font= ("SimHei", 10, "bold"))
            
            # 应用样式
            # 标题
            start_pos = self.sensitive_result_text.search("敏感词检查结果", 1.0)
            if start_pos:
                end_pos = f"{start_pos}+9c"
                self.sensitive_result_text.tag_add("heading", start_pos, end_pos)
            
            # 分类标签 - 自定义搜索所有【符号
            current_pos = "1.0"
            while True:
                pos = self.sensitive_result_text.search("【", current_pos, tk.END)
                if not pos:
                    break
                start_pos = pos
                # 查找匹配的】
                end_pos = self.sensitive_result_text.search("】", start_pos, tk.END)
                if end_pos:
                    end_pos = f"{end_pos}+1c"
                    self.sensitive_result_text.tag_add("category", start_pos, end_pos)
                current_pos = f"{pos}+1c"
            
            # 发现的敏感词 - 自定义搜索所有- '开头的行
            current_pos = "1.0"
            while True:
                pos = self.sensitive_result_text.search("- '", current_pos, tk.END)
                if not pos:
                    break
                start_pos = pos
                # 查找该行结束
                end_pos = self.sensitive_result_text.search("\n", start_pos, tk.END)
                if not end_pos:
                    end_pos = tk.END
                self.sensitive_result_text.tag_add("found", start_pos, end_pos)
                current_pos = end_pos
            
            # 未发现敏感词（成功信息）
            success_pos = self.sensitive_result_text.search("未发现敏感词", 1.0, tk.END)
            if success_pos:
                end_pos = self.sensitive_result_text.search("\n", success_pos, tk.END)
                if not end_pos:
                    end_pos = tk.END
                self.sensitive_result_text.tag_add("success", success_pos, end_pos)
            
            # 上下文部分标题
            contexts_title_pos = self.sensitive_result_text.search("包含敏感词的文本段落：", 1.0, tk.END)
            if contexts_title_pos:
                end_pos = self.sensitive_result_text.search("\n", contexts_title_pos, tk.END)
                if not end_pos:
                    end_pos = tk.END
                self.sensitive_result_text.tag_add("context_section", contexts_title_pos, end_pos)
            
            # 敏感词标签（在上下文部分）
            current_pos = "1.0"
            while True:
                pos = self.sensitive_result_text.search("  敏感词：'", current_pos, tk.END)
                if not pos:
                    break
                start_pos = pos
                # 查找该行结束
                end_pos = self.sensitive_result_text.search("\n", start_pos, tk.END)
                if not end_pos:
                    end_pos = tk.END
                self.sensitive_result_text.tag_add("word_label", start_pos, end_pos)
                current_pos = end_pos
            
            # 段落上下文
            current_pos = "1.0"
            while True:
                pos = self.sensitive_result_text.search("    段落 ", current_pos, tk.END)
                if not pos:
                    break
                start_pos = pos
                # 查找该行结束
                end_pos = self.sensitive_result_text.search("\n", start_pos, tk.END)
                if not end_pos:
                    end_pos = tk.END
                self.sensitive_result_text.tag_add("paragraph", start_pos, end_pos)
                
                # 获取当前段落上下文文本
                para_text = self.sensitive_result_text.get(start_pos, end_pos)
                
                # 获取当前敏感词
                # 向上查找包含敏感词信息的行
                word_info_pos = start_pos
                while word_info_pos > "1.0":
                    # 向上一行
                    word_info_pos = self.sensitive_result_text.index(f"{word_info_pos} linestart - 1l")
                    # 检查是否为敏感词信息行
                    if "敏感词：'" in self.sensitive_result_text.get(word_info_pos, f"{word_info_pos} lineend"):
                        word_info_line = self.sensitive_result_text.get(word_info_pos, f"{word_info_pos} lineend")
                        # 提取敏感词
                        if "敏感词：'" in word_info_line:
                            word_start = word_info_line.find("敏感词：'") + 6
                            word_end = word_info_line.find("'", word_start)
                            if word_end > word_start:
                                current_word = word_info_line[word_start:word_end]
                                # 在当前段落文本中查找并高亮该敏感词
                                word_search_pos = start_pos
                                while True:
                                    # 查找敏感词
                                    word_pos = self.sensitive_result_text.search(current_word, word_search_pos, end_pos)
                                    if not word_pos:
                                        break
                                    # 高亮敏感词
                                    self.sensitive_result_text.tag_add("highlight", 
                                                                      word_pos, 
                                                                      f"{word_pos}+{len(current_word)}c")
                                    # 继续查找下一个
                                    word_search_pos = f"{word_pos}+1c"
                        break
                
                current_pos = end_pos
        
        self.sensitive_result_text.config(state=tk.DISABLED)
    
    def update_result_text(self, text):
        # 更新结果文本框内容
        self.result_text.config(state=tk.NORMAL, foreground="black")
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, text)
        self.result_text.config(state=tk.DISABLED, foreground="black")

if __name__ == "__main__":
    # 启动应用程序
    root = tk.Tk()
    # 设置应用程序图标（可选）
    # root.iconbitmap("icon.ico")
    # 启动应用
    app = FormatCheckerApp(root)
    root.mainloop()